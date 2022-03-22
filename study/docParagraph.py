# !/usr/bin/python
# -*- coding:utf-8 -*-
# @author   : GaiusPluto
# @time     : 2022/3/18 11:15
from docx import Document
# 对齐
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()
paragraph = document.add_paragraph("请在此输入你的段落"*10)
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 缩进
from docx.shared import Inches
paragraph_format = paragraph.paragraph_format
paragraph_format.first_line_indent = Inches(0.25)

# 制表位
tab_stops = paragraph_format.tab_stops
tab_stop = tab_stops.add_tab_stop(Inches(1.5))

# from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
# tab_stop = tab_stops.add_tab_stop(Inches(1.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

# 段落间距
paragraph_format.space_before = Inches(0.25)
paragraph_format.space_after = Inches(0.35)

# 行距
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
paragraph_format.line_spacing = Pt(18)
paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

# 分页属性
paragraph_format.keep_with_next = True

# 应用字符格式
# from docx.enum.text import WD_UNDERLINE
run = document.add_paragraph("请在此输入您的正文"*10).add_run()
font = run.font
font.name = "Microsoft YaHei"
font.size = Pt(10)
font.bold = True
font.italic = False
font.underline = True
# or
# font.underline = WD_UNDERLINE.DOT_DASH
# 字体颜色
from docx.shared import RGBColor
font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
# from docx.enum.dml import MSO_THEME_COLOR_INDEX
# font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1

document.save("test_paragraph.docx")

