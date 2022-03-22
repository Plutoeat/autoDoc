# !/usr/bin/python
# -*- coding:utf-8 -*-
# @author   : GaiusPluto
# @time     : 2022/3/18 20:31
from docx import Document

# 访问样式
from docx.enum.style import WD_STYLE_TYPE

document = Document()
styles = document.styles
paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
for style in paragraph_styles:
    print(style.name)

# 应用样式
paragraph = document.add_paragraph("a")
paragraph.style = document.styles['Heading 1']
# or
paragraph = document.add_paragraph("b")
paragraph.style = "List Bullet"
# or
paragraph = document.add_paragraph("c", style='Body Text')
body_text_style = document.styles['Body Text']
document.add_paragraph("d", style=body_text_style)

# 添加或删除样式
style = styles.add_style('MYSELFSTYLE', WD_STYLE_TYPE.PARAGRAPH)
style.base_style = styles['Normal']
print(len(styles))
styles['MYSELFSTYLE'].delete()
print(len(styles))

# 定义字符格式
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE

style = styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(14)
font.bold = True
font.italic = False
# font.italic = None
font.underline = WD_UNDERLINE.DOT_DASH

# 定义段落格式
from docx.shared import Inches

document = Document()
style = document.styles.add_style('Indent', WD_STYLE_TYPE.PARAGRAPH)
paragraph_format = style.paragraph_format
paragraph_format.left_indent = Inches(0.25)
paragraph_format.first_line_indent = Inches(-0.25)
paragraph_format.space_before = Pt(12)
paragraph_format.widow_control = True
# next_paragraph_style指定要应用于在该样式的段落之后插入的新段落的样式
styles['Heading 1'].next_paragraph_style = styles['Body Text']

document.save("test_style.docx")
