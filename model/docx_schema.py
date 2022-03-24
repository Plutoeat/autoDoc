# !/usr/bin/python
# -*- coding:utf-8 -*-
# @author   : GaiusPluto
# @time     : 2022/3/19 16:54
import yaml
from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from model.getAlignment import get_paragraph_alignment
from model.getFontSize import get_font_size
from model.getSpaceRule import get_space_rule


class DocxSchema:
    def __init__(self, title, **kwargs):
        """
        All rules come from Dalian University
        :param title: doc's title
        """
        self.title = title
        self.document = Document()
        self.sections = self.document.sections
        with open('./model/config.yaml', 'r', encoding='utf-8') as f:
            self.data = yaml.load(f.read(), Loader=yaml.FullLoader)
            f.close()

    def save(self):
        self.document.save("%s.docx" % self.title)


class Paper(DocxSchema):
    def __init__(self, title):
        super().__init__(title)
        self.docx_model = self.data['model'][0]
        self.sections[0].page_width = Cm(self.docx_model['page']['width'])
        self.sections[0].page_height = Cm(self.docx_model['page']['height'])
        # 默认横向，有特殊需求则为纵向
        if self.docx_model['page']['direction']:
            self.sections[0].orientation = WD_ORIENTATION.LANDSCAPE
        # 页面设置
        # 设置页眉
        self.document.sections[0].header_distance = Cm(self.docx_model['page']['header_distance'])
        # 设置页脚，页码需要自行设置
        self.document.sections[0].footer_distance = Cm(self.docx_model['page']['footer_distance'])
        # 页边距一律采取
        self.sections[0].top_margin = Cm(self.docx_model['page']['top_margin'])
        self.sections[0].bottom_margin = Cm(self.docx_model['page']['bottom_margin'])
        self.sections[0].left_margin = Cm(self.docx_model['page']['left_margin'])
        self.sections[0].right_margin = Cm(self.docx_model['page']['right_margin'])
        # 设置样式
        for font_model in self.docx_model['font_models']:
            self.set_style(
                method=font_model['method'],
                style_name=font_model['name'],
                font_name=font_model['font']['CN'],
                font_size=font_model['font']['size'],
                base_style=font_model['base_style'],
                font_west_name=font_model['font']['EN'],
                font_bold=font_model['font']['bold'],
                font_italic=font_model['font']['italic'],
                font_underline=font_model['font']['underline'],
                font_color_r=0,
                font_color_g=0,
                font_color_b=0,
                left_indent=font_model['paragraph']['left_indent'],
                right_indent=font_model['paragraph']['right_indent'],
                first_line_indent=font_model['paragraph']['first_line_indent'],
                space_before=font_model['paragraph']['space_before'],
                space_after=font_model['paragraph']['space_after'],
                line_spacing_rule=font_model['paragraph']['line_spacing_rule'],
                line_spacing=font_model['paragraph']['line_spacing']
            )

        # 中文摘要
        self.document.add_paragraph('摘    要', style='摘要').alignment = get_paragraph_alignment(self.docx_model['font_models'][0]['paragraph']['alignment'])
        self.document.add_paragraph('请在此输入你的正文', style='Normal')
        self.document.add_paragraph('关键词：关键词1；关键词2；关键词3', style='关键词').alignment = get_paragraph_alignment(self.docx_model['font_models'][5]['paragraph']['alignment'])

        # 英文摘要
        self.document.add_section(WD_SECTION_START.ODD_PAGE)
        self.document.add_paragraph('Abstract', style='摘要').alignment = get_paragraph_alignment(self.docx_model['font_models'][0]['paragraph']['alignment'])
        self.document.add_paragraph('Please enter your text here', style='Normal')
        self.document.add_paragraph('Key Words：Key Words1；Key Words2；Key Words3',
                                    style='关键词').alignment = get_paragraph_alignment(self.docx_model['font_models'][5]['paragraph']['alignment'])

        # 添加目录页
        self.document.add_section(WD_SECTION_START.ODD_PAGE)
        self.document.add_paragraph('目    录', style='摘要').alignment = get_paragraph_alignment(self.docx_model['font_models'][0]['paragraph']['alignment'])

        # 添加绪论
        self.document.add_section(WD_SECTION_START.ODD_PAGE)
        self.document.add_paragraph('绪    论', style='摘要').alignment = get_paragraph_alignment(self.docx_model['font_models'][0]['paragraph']['alignment'])
        self.document.add_paragraph('请在此输入你的绪论', style='Normal')
        self.document.add_paragraph('Please enter your introduction here', style='Normal')

        # 添加标题及正文
        self.document.add_section(WD_SECTION_START.ODD_PAGE)
        self.document.add_paragraph('1 请在此输入你的一级标题', style='论文标题 1')
        self.document.add_paragraph('1.1 请在此输入你的二级标题', style='论文标题 2')
        self.document.add_paragraph('1.1.1 请在此输入你的三级标题', style='论文标题 3')
        self.document.add_paragraph('请在此输入你的论文正文', style='Normal')

        # 添加结论
        self.document.add_section(WD_SECTION_START.ODD_PAGE)
        self.document.add_paragraph('结    论', style='论文标题 1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph_format = self.document.add_paragraph('请在此输入你的结论正文', style='Normal').paragraph_format
        paragraph_format.line_spacing = 1.25
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        # 添加参考文献
        self.document.add_section(WD_SECTION_START.ODD_PAGE)
        paragraph = self.document.add_paragraph('', style='论文标题 1')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run('参 考 文 献')
        run.font.size = Pt(12)
        paragraph_format = self.document.add_paragraph('请在此填写参考文献的著录').paragraph_format
        paragraph_format.line_spacing = 1.25
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        # 添加附录
        self.document.add_section(WD_SECTION_START.ODD_PAGE)
        paragraph = self.document.add_paragraph('', style='论文标题 1')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run('附录1 附录内容')
        run.font.size = Pt(15)
        paragraph_format = self.document.add_paragraph('请在此输入附录正文').paragraph_format
        paragraph_format.line_spacing = 1.3
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

    def set_style(
            self,
            method: str,
            style_name: str,
            font_name: str,
            font_size: int,
            base_style: str = None,
            font_west_name: str = None,
            font_bold: bool = False,
            font_italic: bool = False,
            font_underline: bool = False,
            font_color_r: int = 0,
            font_color_g: int = 0,
            font_color_b: int = 0,
            left_indent: int = 0,
            right_indent: int = 0,
            first_line_indent: bool = False,
            space_before: int = 0,
            space_after: int = 0,
            line_spacing_rule: str = '单倍行距',
            line_spacing: bool or int = False
    ) -> None:
        if method == 'add':
            style = self.document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            if base_style is None:
                base_style = 'Normal'
            style.base_style = self.document.styles[base_style]
            style.hidden = False
            style.quick_style = True
        else:
            style = self.document.styles[style_name]
        # 配置字体
        if font_west_name is None or font_west_name == '':
            style.font.name = font_name
        else:
            style.font.name = font_west_name
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        style.font.size = Pt(get_font_size(font_size))
        style.font.bold = font_bold
        style.font.italic = font_italic
        style.font.underline = font_underline
        style.font.color.rgb = RGBColor(font_color_r, font_color_g, font_color_b)
        # 设置段落
        style.paragraph_format.left_indent = Pt(left_indent)
        style.paragraph_format.right_indent = Pt(right_indent)
        if first_line_indent:
            style.paragraph_format.first_line_indent = Pt(get_font_size(font_size) * 2)
        else:
            style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        style.paragraph_format.line_spacing_rule = get_space_rule(line_spacing_rule)
        if line_spacing:
            if line_spacing_rule == '多倍行距':
                style.paragraph_format.line_spacing = line_spacing
            else:
                style.paragraph_format.line_spacing = Pt(line_spacing)
